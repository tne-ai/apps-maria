import json
from io import BytesIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from tne.TNE import TNE

# Initialize the TNE object
session = TNE(uid=UID, bucket_name=BUCKET, project=PROJECT, version=VERSION)

def currency_formatter(value, pos):
    """Format the number with commas and prepend a dollar sign."""
    return '${:,.2f}'.format(value)

def comma_formatter(value, pos):
    """Format the number with commas only, no decimal places."""
    return '{:,.0f}'.format(value)

def format_axis_ticks(ax, callback):
    """
    Apply formatting to the axis ticks based on the callback string.
    We look for patterns in the callback function. For example:
      - "return '$' + value.toLocaleString();" => currency formatting
      - "return value.toLocaleString();" => comma formatting
    """
    if callback and "toLocaleString" in callback:
        if "$" in callback:
            ax.yaxis.set_major_formatter(FuncFormatter(currency_formatter))
        else:
            ax.yaxis.set_major_formatter(FuncFormatter(comma_formatter))
    else:
        # Default formatting if no recognizable callback
        ax.yaxis.set_major_formatter(FuncFormatter(lambda x, pos: f"{x:,.0f}"))

def parse_chart_data(content_data):
    """
    Attempt to parse the chart data which may sometimes be double-encoded 
    or contain additional formatting. We try multiple attempts to decode 
    JSON content.
    """
    try:
        # First try to load as is
        return json.loads(content_data)
    except json.JSONDecodeError:
        # If that fails, the data might be JSON within a string => try loading twice
        try:
            return json.loads(json.loads(content_data))
        except:
            return None

def convert_to_docx(content, output_file):
    # Create a Word document
    doc = Document()

    # ----------------------
    # 1. Header & Footer
    # ----------------------
    section = doc.sections[0]
    if "header_text" in content:
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = content["header_text"]

    if "footer_text" in content:
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = content["footer_text"]

    # ----------------------
    # 2. Document Title
    # ----------------------
    if "doc_title" in content:
        doc.add_heading(content["doc_title"], level=0)

    # ----------------------
    # 3. Process Sections
    # ----------------------
    for section_data in content.get("sections", []):
        content_type = section_data.get("type", "").lower()
        content_text = section_data.get("content", "")

        # (A) Heading for the section
        if "heading" in section_data:
            # Adjust level as needed (1,2,3,...)
            doc.add_heading(section_data["heading"], level=2)

        # (B) Handle content by type
        if content_type == "raw text":
            paragraph = doc.add_paragraph(content_text)
            # Apply optional style
            if "style" in section_data:
                paragraph.style = section_data["style"]

        elif content_type == "table":
            lines = content_text.strip().split("\n")
            headers = lines[0].split("|")
            rows = [line.split("|") for line in lines[1:]]

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'

            # (i) Headers
            hdr_cells = table.rows[0].cells
            for i, header_text in enumerate(headers):
                hdr_cells[i].text = header_text.strip()

            # (ii) Rows
            for row_data in rows:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data.strip()

            # (iii) Caption if present
            if "caption" in section_data:
                caption_para = doc.add_paragraph(section_data["caption"])
                caption_para.style = "Caption"

        elif content_type == "chart":
            chart_info = parse_chart_data(content_text)
            if chart_info is None:
                # Could not parse chart data, skip
                doc.add_paragraph("ERROR: Chart data could not be parsed.")
                continue

            try:
                # Advanced chart rendering
                chart_type = chart_info.get("type", "line")
                data = chart_info.get("data", {})
                options = chart_info.get("options", {})

                labels = data.get("labels", [])
                datasets = data.get("datasets", [])

                # Create figure and main axis
                fig, ax = plt.subplots(figsize=(6, 4))

                # Handle multiple y-axes from options
                y_axes_config = options.get("scales", {}).get("yAxes", [])
                axis_map = {}

                # If there's at least one y-axis config, use it as primary
                if y_axes_config:
                    primary_yaxis_conf = y_axes_config[0]
                    axis_map[primary_yaxis_conf.get("id", "y-axis-0")] = ax

                    # Additional Y-axes if any
                    for extra_yaxis_conf in y_axes_config[1:]:
                        twin_ax = ax.twinx()
                        axis_map[extra_yaxis_conf.get("id", "y-axis-1")] = twin_ax
                else:
                    # No config => single default axis
                    axis_map["y-axis-0"] = ax

                # Configure each y-axis
                for yaxis_conf in y_axes_config:
                    y_id = yaxis_conf.get("id", "y-axis-0")
                    current_ax = axis_map[y_id]

                    # Set axis label
                    scale_label = yaxis_conf.get("scaleLabel", {})
                    if scale_label.get("display", False):
                        current_ax.set_ylabel(scale_label.get("labelString", ""), fontsize=10)

                    # Tick formatting
                    ticks_conf = yaxis_conf.get("ticks", {})
                    callback = ticks_conf.get("callback", "")
                    format_axis_ticks(current_ax, callback)

                # If no custom config, apply a default formatter
                if not y_axes_config:
                    ax.yaxis.set_major_formatter(FuncFormatter(lambda x, pos: f"{x:,.0f}"))

                # X-axis label if provided
                x_axes_config = options.get("scales", {}).get("xAxes", [])
                if x_axes_config:
                    x_axis_conf = x_axes_config[0]
                    scale_label = x_axis_conf.get("scaleLabel", {})
                    if scale_label.get("display", False):
                        ax.set_xlabel(scale_label.get("labelString", ""), fontsize=10)
                else:
                    ax.set_xlabel("")

                # Plot each dataset on the appropriate axis
                for dataset in datasets:
                    yAxisID = dataset.get("yAxisID", "y-axis-0")
                    plot_ax = axis_map.get(yAxisID, ax)

                    color = dataset.get("borderColor", "#000")
                    label = dataset.get("label", "")
                    line_data = dataset.get("data", [])

                    if chart_type == "line":
                        plot_ax.plot(labels, line_data, label=label, color=color, marker="o")
                    elif chart_type == "bar":
                        plot_ax.bar(labels, line_data, label=label, color=color)
                    else:
                        # Fallback to line if not recognized
                        plot_ax.plot(labels, line_data, label=label, color=color, marker="o")

                # Title
                title_conf = options.get("title", {})
                if title_conf.get("display", False):
                    plt.title(title_conf.get("text", ""), fontsize=12)

                # Grid
                ax.grid(True)

                # Legend
                legend_conf = options.get("legend", {})
                if legend_conf.get("display", True):
                    # Combine legends from all axes
                    handles, labels_legend = [], []
                    for axis_id, axis_obj in axis_map.items():
                        h, l = axis_obj.get_legend_handles_labels()
                        handles.extend(h)
                        labels_legend.extend(l)
                    if handles:
                        ax.legend(handles, labels_legend, loc='best')

                # Save to buffer
                chart_stream = BytesIO()
                plt.savefig(chart_stream, format='png')
                plt.close(fig)
                chart_stream.seek(0)

                # Insert chart
                doc.add_picture(chart_stream, width=Inches(5.5))
                chart_stream.close()

                # Add caption if present
                if "caption" in section_data:
                    caption_para = doc.add_paragraph(section_data["caption"])
                    caption_para.style = "Caption"

            except Exception as e:
                doc.add_paragraph(f"ERROR GENERATING CHART: {e}")
                continue

        else:
            # Unrecognized type
            doc.add_paragraph(f"(Unrecognized content type: {content_type})")

    # ----------------------
    # 4. Save & Upload
    # ----------------------
    doc.save(output_file)
    session.upload_object(output_file, doc)
    return output_file

# Load JSON content from the input
try:
    content_json = json.loads(PROCESS_INPUT)  # Parse PROCESS_INPUT string into a dictionary
    output_file = content_json["document_filename"]

    # Generate the docx file
    result = convert_to_docx(content_json, output_file)
except json.JSONDecodeError as e:
    result = f"Invalid JSON input: {e}"
