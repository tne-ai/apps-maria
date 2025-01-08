import json
from io import BytesIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
import re
from tne.TNE import TNE

# Initialize the TNE object
session = TNE(uid=UID, bucket_name=BUCKET, project=PROJECT, version=VERSION)

def currency_formatter(value, pos):
    """Format the number with commas and prepend a dollar sign."""
    return '${:,.2f}'.format(value)

def chartjs_color_to_mpl(color_str):
    """
    Convert a Chart.js-style RGBA color string to Matplotlib format.
    """
    pattern = r'rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)(?:\s*,\s*([\d\.]+))?\)'
    match = re.match(pattern, color_str.strip())
    if match:
        r, g, b, a = match.groups()
        return (int(r) / 255.0, int(g) / 255.0, int(b) / 255.0, float(a) if a else 1.0)
    return color_str  # Return as-is for other valid formats.

def parse_chart_data(content_data):
    """
    Parse chart data JSON.
    """
    try:
        return json.loads(content_data)
    except json.JSONDecodeError:
        return None

def convert_to_docx(content, output_file):
    # Create a Word document
    doc = Document()

    # Add header and footer if provided
    section = doc.sections[0]
    if "header_text" in content:
        section.header.paragraphs[0].text = content["header_text"]
    if "footer_text" in content:
        section.footer.paragraphs[0].text = content["footer_text"]

    # Add document title
    if "doc_title" in content:
        doc.add_heading(content["doc_title"], level=0)

    # Process sections
    for section_data in content.get("sections", []):
        # Add heading
        if "heading" in section_data:
            doc.add_heading(section_data["heading"], level=2)

        content_type = section_data.get("type", "").lower()
        content_text = section_data.get("content", "")

        if content_type == "raw text":
            doc.add_paragraph(content_text)

        elif content_type == "table":
            rows = [row.split("|") for row in content_text.strip().split("\n")]
            table = doc.add_table(rows=1, cols=len(rows[0]))
            table.style = 'Table Grid'

            # Add headers
            for i, header in enumerate(rows[0]):
                table.rows[0].cells[i].text = header.strip()

            # Add rows
            for row_data in rows[1:]:
                row_cells = table.add_row().cells
                for j, cell_data in enumerate(row_data):
                    row_cells[j].text = cell_data.strip()

        elif content_type == "chart":
            chart_info = parse_chart_data(content_text)
            if not chart_info:
                doc.add_paragraph("ERROR: Chart data could not be parsed.")
                continue

            # Inside the "chart" content_type section
            try:
                # Extract chart data
                labels = chart_info["data"]["labels"]
                datasets = chart_info["data"]["datasets"]
                title = chart_info["options"]["plugins"]["title"].get("text", "")
                ylabel = chart_info["options"]["scales"]["y"]["title"]["text"]
                xlabel = chart_info["options"]["scales"]["x"]["title"]["text"]
            
                # Create the plot
                fig, ax = plt.subplots(figsize=(6, 4))
                for dataset in datasets:
                    color = chartjs_color_to_mpl(dataset["borderColor"])
                    ax.plot(labels, dataset["data"], label=dataset["label"], color=color, marker="o")
            
                # Configure the plot
                ax.set_title(title)
                ax.set_ylabel(ylabel)
                ax.set_xlabel(xlabel)
                ax.grid(True)
                ax.legend()
            
                # Use tight layout to ensure nothing is cut off
                plt.tight_layout()
            
                # Save to buffer
                chart_stream = BytesIO()
                plt.savefig(chart_stream, format='png')
                plt.close(fig)
                chart_stream.seek(0)
            
                # Insert chart into document
                doc.add_picture(chart_stream, width=Inches(5.5))
                chart_stream.close()
            
            except Exception as e:
                doc.add_paragraph(f"ERROR GENERATING CHART: {e}")

    # Save document
    doc.save(output_file)
    session.upload_object(output_file, doc)
    return output_file

# Load JSON content and generate the document
try:
    content_json = json.loads(PROCESS_INPUT)
    output_file = content_json["document_filename"]
    result = convert_to_docx(content_json, output_file)
except json.JSONDecodeError as e:
    result = f"Invalid JSON input: {e}"
