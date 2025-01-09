import json
from io import BytesIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from tne.TNE import TNE
import re

# Initialize the TNE object
session = TNE(uid=UID, bucket_name=BUCKET, project=PROJECT, version=VERSION)

def chartjs_color_to_mpl(color_str):
    """
    Convert a Chart.js-style RGBA color string to Matplotlib format.
    """
    pattern = r'rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)(?:\s*,\s*([\d\.]+))?\)'
    match = re.match(pattern, color_str.strip())
    if match:
        r, g, b, a = match.groups()
        return (int(r) / 255.0, int(g) / 255.0, int(b) / 255.0, float(a) if a else 1.0)
    return color_str  # Return as-is for other valid formats.

def convert_to_docx(content, output_file):
    # Create a Word document
    doc = Document()

    # -- 1. Set up header and footer (if provided) --
    section = doc.sections[0]
    if "header_text" in content:
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = content["header_text"]

    if "footer_text" in content:
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = content["footer_text"]

    # -- 2. Add a main title to the document (if provided) --
    if "doc_title" in content:
        doc.add_heading(content["doc_title"], level=0)

    # -- 3. Process each section in the content --
    for section_data in content["sections"]:
        content_type = section_data["type"]
        content_text = section_data["content"]

        # (A) Optionally add a heading for this section
        if "heading" in section_data:
            doc.add_heading(section_data["heading"], level=2)

        # (B) Check the content type and handle accordingly
        if content_type == "raw text":
            paragraph = doc.add_paragraph(content_text)
            if "style" in section_data:
                paragraph.style = section_data["style"]

        elif content_type == "table":
            lines = content_text.strip().split("\n")
            headers = lines[0].split("|")
            rows = [line.split("|") for line in lines[1:]]

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'

            # Add header row
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header.strip()

            # Add data rows
            for row_data in rows:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data.strip()

            # Add caption if present
            if "caption" in section_data:
                caption_para = doc.add_paragraph(section_data["caption"])
                caption_para.style = "Caption"

        elif content_type == "chart":
            try:
                chart_info = json.loads(content_text)

                plt.figure(figsize=(6, 4))
                for dataset in chart_info["data"]["datasets"]:
                    # Use chartjs_color_to_mpl for parsing color strings
                    color = chartjs_color_to_mpl(dataset.get("borderColor", "#000"))
                    plt.plot(chart_info["data"]["labels"],
                             dataset["data"],
                             label=dataset["label"],
                             color=color,
                             marker="o")

                plt.title(chart_info["options"]["title"]["text"])
                plt.xlabel('Year')
                plt.ylabel('Value')
                plt.grid(True)
                if chart_info["options"]["legend"]["display"]:
                    plt.legend()

                # Use tight layout to ensure nothing is cut off
                plt.tight_layout()

                # Save chart to a BytesIO buffer
                chart_stream = BytesIO()
                plt.savefig(chart_stream, format='png')
                plt.close()
                chart_stream.seek(0)

                # Insert chart image into the document
                doc.add_picture(chart_stream, width=Inches(5.5))
                chart_stream.close()

                # Add caption if present
                if "caption" in section_data:
                    caption_para = doc.add_paragraph(section_data["caption"])
                    caption_para.style = "Caption"

            except Exception as e:
                paragraph = doc.add_paragraph(f"ERROR GENERATING CHART: {e}")
                continue

    # -- 4. Save and upload the document --
    doc.save(output_file)
    session.upload_object(output_file, doc)
    return output_file

# Load JSON content from the input
try:
    content_json = json.loads(PROCESS_INPUT)  # Parse PROCESS_INPUT string into a dictionary
    output_file = content_json["document_filename"]

    # Generate the docx file
    result = convert_to_docx(content_json, output_file)
except json.JSONDecodeError as e:
    result = f"Invalid JSON input: {e}"
